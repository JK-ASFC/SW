# ─── Génération du guide Burger Quiz en Word ───────────────────────────────
# Prérequis : pip install python-docx
# Lancement  : python generate_guide.py
# Résultat   : Guide_BurgerQuiz.docx dans le même dossier
# ───────────────────────────────────────────────────────────────────────────

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

# ── Vérification de la librairie ──────────────────────────────────────────
try:
    from docx import Document
except ImportError:
    print("\n[ERREUR] La librairie python-docx n'est pas installée.")
    print("Lancez cette commande puis relancez le script :\n")
    print("    pip install python-docx\n")
    sys.exit(1)

# ── Couleurs ──────────────────────────────────────────────────────────────
JAUNE   = RGBColor(0xED, 0xD4, 0x00)
ROUGE   = RGBColor(0xCC, 0x00, 0x00)
NOIR    = RGBColor(0x11, 0x11, 0x11)
GRIS    = RGBColor(0x44, 0x44, 0x44)
BLANC   = RGBColor(0xFF, 0xFF, 0xFF)
VERT    = RGBColor(0x05, 0x96, 0x69)
BLEU    = RGBColor(0x08, 0x91, 0xB2)
BG_DARK = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()

# ── Marges de page ────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles de base ────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(10.5)
style_normal.font.color.rgb = NOIR


def set_cell_bg(cell, hex_color):
    """Colore le fond d'une cellule de tableau."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def h1(text):
    """Titre principal de section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(22)
    run.font.color.rgb = ROUGE
    # Ligne de séparation
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'),  '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'EDD400')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(text):
    """Sous-titre."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xED, 0xD4, 0x00)
    return p


def h3(text):
    """Sous-sous-titre."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS
    return p


def body(text):
    """Paragraphe normal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def bullet(text, bold_prefix=None):
    """Item de liste."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10.5)
        r  = p.add_run(text)
        r.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
    return p


def callout(label, text, color_hex='EDD400'):
    """Boîte d'avertissement / info."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_bg(cell, '1A1A1A')
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    rb = p.add_run(label + ' ')
    rb.bold = True
    rb.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16)
    )
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    r.font.size = Pt(10)
    doc.add_paragraph()  # espace après


def key_table(headers, rows):
    """Tableau de raccourcis clavier."""
    col_count = len(headers)
    t = doc.add_table(rows=1, cols=col_count)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # En-tête
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], '1A1A1A')
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h.upper())
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Lignes
    for row_data in rows:
        row_cells = t.add_row().cells
        for i, cell_text in enumerate(row_data):
            p = row_cells[i].paragraphs[0]
            if i == 0:
                r = p.add_run(cell_text)
                r.bold = True
                r.font.name = 'Courier New'
                r.font.size = Pt(10)
                r.font.color.rgb = JAUNE
            else:
                r = p.add_run(cell_text)
                r.font.size = Pt(10)
    doc.add_paragraph()


def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Cm(3)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BURGER QUIZ')
r.bold      = True
r.font.size = Pt(40)
r.font.color.rgb = JAUNE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('GUIDE D\'UTILISATION')
r.bold      = True
r.font.size = Pt(20)
r.font.color.rgb = ROUGE

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Application d\'animation de quiz — Usage local')
r.font.size = Pt(12)
r.font.color.rgb = GRIS

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Conception & Développement : JK')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Version 2 · 2025')
r.font.size = Pt(10)
r.font.color.rgb = GRIS

page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  1. PRÉREQUIS
# ═══════════════════════════════════════════════════════════════════════════
h1('1. Prérequis')

body('Burger Quiz est une application web locale. Elle fonctionne dans un navigateur moderne sans installation, '
     'mais nécessite un serveur HTTP local pour que les fichiers se chargent correctement (vidéos, sons, images).')

h2('Option 1 — Python (recommandé)')
body('Python est gratuit et disponible sur python.org. Une fois installé, ouvrez un terminal dans le dossier '
     'du projet et lancez :')

p = doc.add_paragraph()
r = p.add_run('    python -m http.server 8000')
r.font.name = 'Courier New'
r.font.size = Pt(11)
r.font.color.rgb = JAUNE

body('Puis ouvrez votre navigateur à l\'adresse : http://localhost:8000')

callout('💡 Raccourci Windows —',
        'Créez un fichier texte contenant "python -m http.server 8000 && pause", '
        'renommez-le start_server.bat, double-cliquez pour lancer le serveur.',
        'EDD400')

h2('Option 2 — Extension VS Code')
body('Installez l\'extension Live Server dans Visual Studio Code, puis cliquez sur "Go Live" '
     'en bas à droite. Le navigateur s\'ouvre automatiquement.')

h2('Option 3 — XAMPP / WAMP')
body('Placez le dossier dans htdocs (XAMPP) ou www (WAMP) et accédez via http://localhost/burger/')

callout('⚠️ Important —',
        'Ne jamais ouvrir index.html en double-cliquant (protocole file://). '
        'Cela bloque le chargement des vidéos et des sons.',
        'CC0000')

h2('Navigateurs compatibles')
bullet('Google Chrome ou Microsoft Edge — recommandés (meilleure compatibilité vidéo)')
bullet('Firefox — compatible, mais certains formats vidéo peuvent poser problème')
bullet('Safari — compatible, nécessite des fichiers MP4 H.264')

h2('Formats de médias supportés')
bullet('Vidéo : MP4 (H.264), WebM')
bullet('Audio : MP3, OGG, WAV, M4A')
bullet('Image : JPG, PNG, GIF, WebP, SVG')
bullet('Document : PDF')

page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  2. LANCEMENT
# ═══════════════════════════════════════════════════════════════════════════
h1('2. Lancement')

body('Une fois le serveur démarré, rendez-vous sur http://localhost:8000. '
     'L\'écran d\'accueil propose trois destinations :')

bullet('⚙ Admin — Préparer et gérer les slides, listes, médias')
bullet('▶ Jouer — Choisir une liste et lancer le quiz')
bullet('📖 Guide — Cette documentation')

callout('ℹ️ Workflow recommandé —',
        'Préparez vos slides en Admin, exportez-les en JSON pour les sauvegarder, '
        'puis lancez le jeu le soir de l\'émission.',
        '0891B2')

page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  3. ADMINISTRATION
# ═══════════════════════════════════════════════════════════════════════════
h1('3. Administration')

body('L\'interface d\'administration est accessible via le bouton Admin de l\'accueil '
     'ou en appuyant sur la touche A pendant le jeu.')

h2('3.1  Gestion des listes')
body('Le widget "Liste ▾" en haut permet de gérer plusieurs listes de questions indépendantes '
     '(ex : "Saison 1", "Spécial Noël"). Chaque liste possède ses propres slides.')
bullet('Cliquer sur le widget pour ouvrir le panneau de gestion')
bullet('+ Liste — créer une nouvelle liste vide')
bullet('Cliquer sur une liste pour la rendre active (étoile ★)')
bullet('Modifier le nom directement dans le champ de texte')
bullet('Supprimer une liste avec le bouton ✕ (les slides sont conservés)')

h2('3.2  Gestion des slides')
bullet('+ Nouveau slide — ajoute un slide à la liste active')
bullet('Cliquer sur un slide dans la liste de gauche pour l\'éditer')
bullet('Glisser-déposer pour réordonner les slides')
bullet('Bouton ✕ rouge sur chaque slide pour le supprimer')
bullet('Les modifications sont sauvegardées automatiquement')

h2('3.3  Types de slides')
key_table(
    ['Type', 'Nom', 'Description'],
    [
        ('🎬  V — Vidéo',       'Vidéo',         'Lecture plein écran. Utile pour l\'intro ou des extraits vidéo.'),
        ('🍟  N — Nuggets',     'Nuggets',        'QCM 4 réponses (A/B/C/D). Bonne réponse révélée avec Espace.'),
        ('🧂  S — Sel/Poivre',  'Sel ou Poivre',  'Vrai/Faux. SEL = vrai, POIVRE = faux.'),
        ('🖼️  I — Image',       'Image',          'Affiche une image avec titre et question optionnels.'),
        ('🍔  M — Menu',        'Menu',           'Question avec sous-questions. Navigation ← → dans les sous-questions.'),
        ('📋  L — Liste',       'Liste menus',    'Liste de 3 menus jouables indépendamment.'),
        ('➕  A — Addition',    'Addition',       'Round addition : éléments à additionner.'),
        ('💀  B — Burger Mort', 'Burger de la Mort', 'Slide déclencheur du BDM. Contient la vidéo d\'intro.'),
        ('👥  C — Candidats',  'Candidats',      'Présentation des deux candidats en deux temps (Espace x2).'),
        ('🎁  G — Cadeau',      'Cadeau',         'Révélation du cadeau BDM affiché si l\'équipe réussit.'),
        ('🎲  T — Toss',        'Toss',           'Tirage au sort / duel. Espace révèle le résultat.'),
    ]
)

h2('3.4  Attacher un média')
body('La plupart des types acceptent un média (image, vidéo, audio, PDF). '
     'Glissez-déposez dans la zone pointillée ou cliquez pour parcourir. '
     'Le fichier est stocké dans le navigateur (IndexedDB).')

h2('3.5  Bibliothèque de médias')
body('Le bouton "🖼 Médias" ouvre la galerie de tous les fichiers stockés. '
     'Vous pouvez y supprimer les médias inutilisés pour libérer de l\'espace.')

page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  4. BURGER DE LA MORT
# ═══════════════════════════════════════════════════════════════════════════
h1('4. Burger de la Mort (Admin)')

body('Le Burger de la Mort est le grand défi final de l\'émission. '
     'Le bouton "💀 Listes BDM" dans l\'administration ouvre le gestionnaire.')

h2('4.1  Créer une liste BDM')
bullet('Cliquer "+ Nouvelle liste"')
bullet('Nommer la liste (ex : "Liste A", "Spéciale")')
bullet('Saisir jusqu\'à 10 questions dans les champs numérotés')
bullet('Cliquer "★ Définir active" pour choisir la liste utilisée au prochain BDM')

callout('💡 Conseil —',
        'Préparez 2 à 3 listes à l\'avance. Si une équipe échoue, vous pouvez '
        'basculer sur une autre liste pour l\'équipe adverse sans rejouer les mêmes questions.',
        'EDD400')

h2('4.2  Slide déclencheur')
body('Créez un slide de type 💀 B dans votre liste principale. Il contient la vidéo d\'intro du BDM. '
     'Quand ce slide apparaît pendant le jeu, appuyez sur → pour lancer la vidéo, '
     'puis → à nouveau pour entrer dans les questions.')

page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  5. PARAMÈTRES
# ═══════════════════════════════════════════════════════════════════════════
h1('5. Paramètres (Admin)')

body('Le bouton "⚙ Paramètres" permet de configurer les médias globaux de l\'émission.')

key_table(
    ['Paramètre', 'Rôle'],
    [
        ('Musique BDM',        'Musique jouée en boucle pendant les questions du Burger de la Mort.'),
        ('Musique de victoire','Musique déclenchée quand une équipe atteint 25 points (dure 11 sec).'),
        ('Générique de fin',   'Vidéo diffusée après le BDM réussi (générique façon émission TV).'),
        ('Retour au calme',    'Titre et message affichés après le générique sur l\'écran sonomètre.'),
    ]
)

page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  6. EXPORT / IMPORT
# ═══════════════════════════════════════════════════════════════════════════
h1('6. Export / Import')

h2('6.1  Exporter')
body('Le bouton "⬇ Exporter" génère un fichier .json contenant :')
bullet('Tous les slides de toutes les listes')
bullet('Toutes les listes (noms, IDs) et listes BDM avec leurs questions')
bullet('Tous les médias (images, vidéos, sons) encodés en base64')
bullet('Les paramètres (musiques, générique, retour au calme)')

callout('⚠️ Sauvegardez régulièrement ! —',
        'Le stockage navigateur peut être effacé si vous videz le cache. '
        'L\'export JSON est votre seule sauvegarde externe.',
        'EDD400')

h2('6.2  Importer')
body('Le bouton "⬆ Importer" restaure un fichier JSON précédemment exporté. '
     'Tous les slides, listes et médias sont rechargés.')

callout('🔴 Attention —',
        'L\'import écrase les données actuelles. Exportez d\'abord si vous '
        'souhaitez conserver la session en cours.',
        'CC0000')

h2('6.3  Import CSV')
body('Il est possible d\'importer une liste de questions au format CSV pour créer '
     'rapidement des slides Nuggets en masse.')

page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  7. TOUCHES CLAVIER — JEU
# ═══════════════════════════════════════════════════════════════════════════
h1('7. Touches clavier — Jeu')

body('Le jeu est entièrement contrôlé au clavier. Voici toutes les touches disponibles.')

h2('7.1  Navigation générale')
key_table(
    ['Touche', 'Action'],
    [
        ('→',     'Slide suivant'),
        ('←',     'Slide précédent'),
        ('P',     'Rejouer la vidéo du slide courant depuis le début'),
        ('R',     'Réinitialiser complètement (scores + position)'),
        ('W',     'Retourner à l\'accueil'),
        ('A',     'Ouvrir l\'administration (nouvel onglet)'),
    ]
)

h2('7.2  Scores')
key_table(
    ['Touche', 'Action'],
    [
        ('M',  '+1 point à l\'équipe Mayo'),
        ('L',  '−1 point à l\'équipe Mayo'),
        ('K',  '+1 point à l\'équipe Ketchup'),
        ('J',  '−1 point à l\'équipe Ketchup'),
    ]
)
body('Quand une équipe atteint 25 points : confettis + musique de victoire (11 secondes) '
     'puis confirmation pour officialiser.')

h2('7.3  Réponses (Nuggets / Sel ou Poivre)')
key_table(
    ['Touche', 'Action'],
    [
        ('1',      'Sélectionner / valider la réponse A'),
        ('2',      'Sélectionner / valider la réponse B'),
        ('3',      'Sélectionner / valider la réponse C'),
        ('4',      'Sélectionner / valider la réponse D'),
        ('Espace', 'Révéler la bonne réponse'),
        ('0',      'Effacer la sélection / masquer la réponse'),
    ]
)

h2('7.4  Mode Menu')
key_table(
    ['Touche', 'Action'],
    [
        ('1 / 2 / 3', 'Entrer dans le sous-menu 1, 2 ou 3'),
        ('→ / ←',     'Naviguer entre les questions du menu'),
        ('Espace',    'Révéler la réponse de la question'),
        ('0',         'Masquer la réponse'),
        ('Échap',     'Quitter le sous-menu, retour au hub'),
    ]
)

h2('7.5  Burger de la Mort')
key_table(
    ['Touche', 'Action'],
    [
        ('→',          'Question suivante (ou lancer les questions après la vidéo)'),
        ('←',          'Question précédente (ou retour aux questions depuis le récap)'),
        ('E / Échap',  'Quitter le BDM'),
        ('M L K J',    'Modifier les scores même pendant le BDM'),
    ]
)

h2('7.6  Générique de fin / Retour au calme')
key_table(
    ['Touche', 'Action'],
    [
        ('→ / Espace', 'Avancer : afficher le retour au calme puis terminer'),
        ('E / Échap',  'Quitter directement'),
    ]
)

page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  8. DÉROULEMENT BDM COMPLET
# ═══════════════════════════════════════════════════════════════════════════
h1('8. Déroulement BDM — Étape par étape')

steps = [
    ('Étape 1', 'Slide 💀 affiché',       'Appuyer → pour lancer la vidéo d\'intro BDM.'),
    ('Étape 2', 'Vidéo intro BDM',         'Laisser la vidéo jouer, puis appuyer → pour entrer dans les questions.'),
    ('Étape 3', '10 questions BDM',        'Navigation ← → entre les questions. La musique BDM joue en boucle.'),
    ('Étape 4', 'Récapitulatif automatique','Après la dernière question, le récap s\'affiche avec toutes les questions.'),
    ('Étape 5', 'RÉUSSI ou ÉCHOUÉ',        'L\'animateur choisit le résultat selon les réponses de l\'équipe.'),
]

for etape, titre, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    rb = p.add_run(etape + ' — ')
    rb.bold = True
    rb.font.color.rgb = ROUGE
    rt = p.add_run(titre + ' : ')
    rt.bold = True
    rt.font.color.rgb = JAUNE
    rd = p.add_run(desc)
    rd.font.color.rgb = GRIS

doc.add_paragraph()

h2('Si RÉUSSI ✅')
bullet('Confettis + musique de joie automatiques')
bullet('Affichage du slide Cadeau (🎁)')
bullet('Bouton "Autre équipe" : relancer un BDM avec une autre liste')
bullet('Bouton "Générique de fin" : lancer la vidéo générique de l\'émission')
bullet('Bouton "Terminer" : fermer la fenêtre BDM')

h2('Si ÉCHOUÉ ❌')
bullet('Un sélecteur de liste apparaît')
bullet('Choisir une autre liste BDM pour l\'équipe adverse')
bullet('Le BDM repart depuis le début avec les nouvelles questions')

h2('Générique de fin')
body('La vidéo configurée en paramètres se lance en plein écran. '
     'Appuyer → ou Espace pour passer en mode Retour au calme.')

h2('Retour au calme (sonomètre)')
body('Écran minimaliste avec un sonomètre en temps réel utilisant le microphone de l\'ordinateur :')
bullet('Vert (≤ 40 dB) — niveau correct, les invités sont calmes')
bullet('Orange (41–60 dB) — légèrement fort')
bullet('Rouge (> 60 dB) — trop fort')
body('Cliquer "Activer le micro" pour démarrer la mesure. '
     'Le navigateur demandera l\'autorisation d\'accès au microphone.')

page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  9. STOCKAGE
# ═══════════════════════════════════════════════════════════════════════════
h1('9. Stockage des données')

body('Burger Quiz utilise deux mécanismes de stockage navigateur, sans aucun serveur de base de données.')

key_table(
    ['Mécanisme', 'Contenu', 'Limite'],
    [
        ('IndexedDB',    'Slides, médias (images, vidéos, sons) en binaire', 'Plusieurs Go selon navigateur'),
        ('LocalStorage', 'Listes, listes BDM, scores, paramètres',           '~5 Mo'),
    ]
)

callout('⚠️ Attention —',
        'Ces données sont liées au navigateur et au domaine. Vider le cache, changer de navigateur '
        'ou accéder depuis une autre machine efface toutes les données. Exportez régulièrement en JSON.',
        'EDD400')

h3('Libérer de l\'espace')
body('Admin → 🖼 Médias → supprimer les fichiers non utilisés.')

page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  10. CRÉDITS
# ═══════════════════════════════════════════════════════════════════════════
h1('10. Crédits')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('JK')
r.bold      = True
r.font.size = Pt(60)
r.font.color.rgb = JAUNE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Conception & Développement')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = GRIS

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Burger Quiz — Application locale')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Inspiré de l\'émission Burger Quiz d\'Alain Chabat.')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Version 2 · 2025 · Tous droits réservés JK')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


# ═══════════════════════════════════════════════════════════════════════════
#  SAUVEGARDE
# ═══════════════════════════════════════════════════════════════════════════
output = 'Guide_BurgerQuiz.docx'
doc.save(output)
print(f'\nFichier genere : {output}\n')
